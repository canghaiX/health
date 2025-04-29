from fastapi import APIRouter,File,UploadFile
from config import get_logger
import torch
import json
from transformers import AutoTokenizer, AutoModel
from pymilvus.model.reranker import BGERerankFunction
from pymilvus.model.hybrid import BGEM3EmbeddingFunction
from pymilvus import (
    connections,
    utility,
    FieldSchema,
    CollectionSchema,
    DataType,
    Collection,
    MilvusException,
    AnnSearchRequest,
    WeightedRanker,
    RRFRanker,
    Function,
    FunctionType
)
from fastapi import FastAPI,Body,HTTPException
import os
from typing import Optional
import pdfplumber
from docx import Document
import shutil
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from routers.file_mineru import process_pdf_file
#from routers.chunk1 import process_markdown_file
from routers.chunk3 import EnhancedMarkdownParser
from routers.chunk7 import split_markdown_into_chunks
from routers.chunk11 import process_markdown_file
from langchain.embeddings.huggingface import HuggingFaceEmbeddings
import requests
import numpy as np
from openai import OpenAI

router = APIRouter()
logger = get_logger(__name__)
client = OpenAI(
    api_key="sk-762b0529add947b081778614c7fe1cda",
    base_url="https://compatible-mode/v1"
)

# 文件存储目录配置
base_dir = "./knowledge_base"
os.makedirs(base_dir, exist_ok=True)
collection_mapping_dir = "./test/collection_mapping.json"
# 准备嵌入模型
# 指定本地模型路径
model_path = "/home/ysdx2025/embedding_model"
#md解析路径
output_dir = './file'
#mineru解析url
url = 'http://localhost:7000/parse-files/'

# 加载模型和分词器
tokenizer = AutoTokenizer.from_pretrained(model_path)
model = AutoModel.from_pretrained(model_path)
model.eval()


# 全局变量，用于存储已经加载好的嵌入模型
embedding_model = None

def load_embedding_model():
    global embedding_model
    if embedding_model is None:
        embedding_model = HuggingFaceEmbeddings(
            model_name=model_path,
            model_kwargs={'device': 'cuda'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
# 定义获取文本嵌入的函数
def get_text_embeddings(texts):
    inputs = tokenizer(texts, padding=True, truncation=True, return_tensors="pt")
    with torch.no_grad():
        outputs = model(**inputs)
        embeddings = torch.mean(outputs.last_hidden_state, dim=1)
        embeddings = torch.nn.functional.normalize(embeddings, p=2, dim=1)
    return embeddings.cpu().numpy()

# 分块函数
def split_text_into_chunks(texts, max_length):
    all_chunks = []
    current_chunk = ""
    for text in texts:
        if len(current_chunk) + len(text) <= max_length:
            # 如果当前块加上当前文本长度不超过最大长度，就合并
            current_chunk += text
        else:
            if current_chunk:
                # 如果当前块不为空，将其添加到结果列表
                all_chunks.append(current_chunk)
            if len(text) <= max_length:
                # 如果当前文本长度不超过最大长度，将其作为新的当前块
                current_chunk = text
            else:
                # 如果当前文本长度超过最大长度，对其进行分块
                for i in range(0, len(text), max_length):
                    part = text[i:i + max_length]
                    all_chunks.append(part)
                current_chunk = ""
    if current_chunk:
        # 处理最后一个未完成的块
        all_chunks.append(current_chunk)
    return all_chunks

def emb_text(texts, is_query):
    """
    对输入的文本列表进行向量化处理。

    参数:
    texts (list): 包含多个文本的列表。

    返回:
    list: 每个文本对应的向量表示列表。
    """
    global embedding_model
    
    if embedding_model is None:
        load_embedding_model()
    
    if is_query:
        result = [embedding_model.embed_query(text) for text in texts]
    
    else:
        result = embedding_model.embed_documents(texts)
    
    return result


def generate_embeddings(docs):
    ef = BGEM3EmbeddingFunction(use_fp16=False, device="cuda:0")
    if isinstance(docs, str):
        docs = [docs]  # 将字符串转换为列表
    return ef(docs)

"""
# 文本向量化
def emb_text(texts,is_query):
    #对输入的文本列表进行向量化处理。

    参数:
    texts (list): 包含多个文本的列表。

    返回:
    list: 每个文本对应的向量表示列表。
    
    embeddings = HuggingFaceEmbeddings(model_name=model_path,
                                       model_kwargs={'device': 'cuda'},
                                        encode_kwargs={'normalize_embeddings':True})
    if is_query:
        return [embeddings.embed_query(texts)]
    else:
        return embeddings.embed_documents(texts)
"""


    # inputs = tokenizer(texts, return_tensors='pt', padding=True, truncation=True)
    # with torch.no_grad():
    #     outputs = model(**inputs)
    # embeddings = outputs.last_hidden_state.mean(dim=5).tolist()
    
    # # 确保每个向量的长度为 1024
    # for i, embedding in enumerate(embeddings):
    #     if len(embedding) != 1024:
    #         raise ValueError(f"向量维度不匹配，期望维度为 1024，实际维度为 {len(embedding)}")
    # return embeddings

#加载或创建知识库映射表
def load_mapping(mapping_path):
    try:
        with open(mapping_path, 'r', encoding='utf-8') as f:
            collection_mapping = json.load(f)
    except FileNotFoundError:
        collection_mapping = {}
    return collection_mapping

#写回并保存知识库映射表
def save_mapping(mapping_path,collection_mapping):
    # 保存映射字典
    with open(mapping_path, 'w', encoding='utf-8') as f:
        json.dump(collection_mapping, f, ensure_ascii=False, indent=4)



def convert_whitespace_to_underscore(filename):
    # 将空白字符替换为下划线
    new_filename = filename.replace(" ", "_")
    return new_filename


def convert_underscore_to_whitespace(filename):
    # 将下划线替换为空白字符
    new_filename = filename.replace("_", " ")
    return new_filename

"""
def generate_valid_uuid():
    #生成首字符为字母或下划线的 UUID4
    #（由于 UUID4 的首字符本质上是十六进制数字，实际有效字符为 a-f 的小写字母）
    
    while True:
        # 生成标准 UUID4 并转换为字符串（去除连字符）
        uid = uuid.uuid4().hex  # 示例: "f47ac10b58cc4372a5670e02b2c3d479"
        first_char = uid[0]
        
        # 检查首字符是否为字母（a-f 的小写形式）
        if first_char.isalpha():
            return uid  # 合法 UUID
"""

# 连接 Milvus
connections.connect(uri="./test/milvus.db")


# 定义字段模式
fields = [
        FieldSchema(
            name="pk", dtype=DataType.VARCHAR, is_primary=True, auto_id=True, max_length=100
        ),
        FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=6000),
        FieldSchema(name="sparse_vector", dtype=DataType.SPARSE_FLOAT_VECTOR),
        FieldSchema(name="dense_vector", dtype=DataType.FLOAT_VECTOR, dim=1024),
        FieldSchema(name="filename", dtype=DataType.VARCHAR, max_length=512),
        FieldSchema(name="fileid", dtype=DataType.VARCHAR, max_length=512)
    ]
schema = CollectionSchema(fields)
"""

bm25_function = Function(
    name='bm25',
    function_type=FunctionType.BM25,
    input_field_names=['text'],
    output_field_names='sparse_bm25'
)

# 定义字段模式
fields = [
        FieldSchema(
            name="pk", dtype=DataType.VARCHAR, is_primary=True, auto_id=True, max_length=100
        ),
        FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=5500),
        FieldSchema(name="sparse_bm25", dtype=DataType.SPARSE_FLOAT_VECTOR),
        FieldSchema(name="dense_vector", dtype=DataType.FLOAT_VECTOR, dim=1024),
        FieldSchema(name="filename", dtype=DataType.VARCHAR, max_length=512),
        FieldSchema(name="fileid", dtype=DataType.VARCHAR, max_length=512)
    ]

schema = CollectionSchema(fields,functions=[bm25_function])
"""
"""

# 定义字段模式
fields = [
        FieldSchema(
        name="pk", dtype=DataType.VARCHAR, is_primary=True, auto_id=True, max_length=100
    ),
    FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=5500),
    FieldSchema(name="vector", dtype=DataType.FLOAT_VECTOR, dim=1024),
    FieldSchema(name="filename", dtype=DataType.VARCHAR, max_length=512),
    FieldSchema(name="fileid", dtype=DataType.VARCHAR, max_length=512)
]
schema = CollectionSchema(fields)  
"""



"""创建知识库接口"""
@router.post("/create_knowledge_base")
async def create_knowledge_base(kbId: str = Body(..., embed=True),kbName: str = Body(..., embed=True)):
    logger.info(f"接收到创建知识库请求: {kbName}")
    #载入映射表
    collection_mapping = load_mapping(collection_mapping_dir)
    #首先判断传入知识库名称是否已经在映射表中存在，若存在则报错
    # ---------- 遍历所有值，检查知识库名称是否已存在 ----------
    # 初始化存在标志
    exists = False
    # 遍历所有键值对的值（即每个集合的信息字典）
    for collection_info in collection_mapping.values():
        # 检查值中是否存在 "kbName" 字段，并比对名称
        if collection_info == kbName :
            exists = True
            break

    if exists:
        error_msg = f"知识库映射表 '{kbName}' 已存在"
        logger.warning(error_msg)
        return {"code": 400, "message": f"知识库 {kbName} 已存在"}

    # ---------- 映射表添加并保存 ----------
    try:

        # 把知识库ID作为知识库名称
        collection_name = "_" + kbId.replace('-', '_')
        # 记录映射关系
        collection_mapping[collection_name] = kbName
        
        # 保存更新后的映射表
        save_mapping(collection_mapping_dir,collection_mapping)
        
        logger.info(f"知识库 '{kbName}' 创建成功，ID: {collection_name}")
    
    except Exception as e:
        logger.error(f"创建知识库失败: {str(e)}", exc_info=True)
        return {"code": 500, "message": f"内部错误: {str(e)}"}
    
    #开始新建知识库
    Knowledge_base_dir = os.path.join(base_dir, kbName)
    
    
    # 检查本地文件夹或Milvus集合是否存在
    if os.path.exists(Knowledge_base_dir) or utility.has_collection(collection_name):
        return {"code": 400, "message": f"知识库 {kbName} 已存在"}
 

    try:
        # 创建本地对应知识库目录结构
        os.makedirs(Knowledge_base_dir, exist_ok=False)
    except FileExistsError:
        return {"code": 400, "message": f"知识库 {kbName} 文件夹已存在"}
    except Exception as e:
        return {"code": 500, "message": f"创建文件夹失败: {str(e)}"}
    
    
    # 创建Milvus集合
    try:
        col = Collection(collection_name, schema, consistency_level="Strong")
        #index_params = {"index_type": "IVF_FLAT", "metric_type": "COSINE", "params": {"nlist": 16}}
        #col.create_index("vector", index_params)
        #col.load()
        
        sparse_index = {"index_type": "SPARSE_INVERTED_INDEX", "metric_type": "IP","params": {"term_threshold": 1, "top_k": 100} }
        col.create_index("sparse_vector", sparse_index)
        #sparse_index = {"index_type": "SPARSE_INVERTED_INDEX", "metric_type": "BM25"}
        #col.create_index("sparse_bm25", sparse_index)
        dense_index = {"index_type": "IVF_FLAT", "metric_type": "IP","params": {"nlist": 1024, "nprobe": 32}}
        col.create_index("dense_vector", dense_index)
        col.load()
        return {"code": 200, "message": "知识库创建成功"}
    
    except Exception as e:
        # 回滚：删除已创建的文件夹
        try:
            import shutil
            shutil.rmtree(Knowledge_base_dir)
        except Exception as cleanup_error:
            logger.error(f"清理文件夹失败: {cleanup_error}")
            return {"code": 500, "message": f"创建知识库失败: {str(e)}"}



"""更改知识库接口"""
@router.post("/update_knowledge_base")
async def update_knowledge_base(kbId: str = Body(..., embed=True),kbName: str = Body(..., embed=True)):
    
    logger.info(f"接收到修改知识库请求:")

    try:
        # 载入映射表
        collection_mapping = load_mapping(collection_mapping_dir)
        collection_name = "_" + kbId.replace('-', '_')
        logger.info(f"{collection_name}")
        #首先判断所给kbId对应的知识库是否存在
        if not utility.has_collection(collection_name):
            return {"code": 400, "message": f"知识库不存在"}

        #接着对映射表进行修改即可
        # 初始化存在标志
        exists = False
        old_kbName=""
    
        #遍历映射表所有键值对的键
        for key in collection_mapping.keys():
            # 找到对应的位置对其值进行覆盖改写
            if key == collection_name :
                exists = True
                old_kbName = collection_mapping[key]
                collection_mapping[key] = kbName
                break
        logger.info(f"{collection_mapping}")
        if exists :
            # 更改原知识库存储本地的文件夹名
            old_folder = base_dir + '/' + old_kbName
            new_folder = base_dir + '/' + kbName

            # 使用os.rename()函数进行重命名
            os.rename(old_folder, new_folder)
            # 保存更新后的映射表
            save_mapping(collection_mapping_dir,collection_mapping)
            return {"code": 200, "message": "知识库更改成功"}
        else :
            return {"code": 500, "message": "知识库更改失败"}
    
    except KeyError:
        return {"code": 200, "message": "知识库不存在"}
    
    except Exception as e:
        # 处理其他可能的异常
        return {"code": 500, "message": f"发生未知错误: {str(e)}"}
    


"""删除知识库接口"""
@router.post("/delete_knowledge_base")
async def delete_knowledge_base(kbId: str = Body(..., embed=True)):

    logger.info(f"接收到删除知识库请求:")

    try:
        # 载入映射表
        collection_mapping = load_mapping(collection_mapping_dir)
        collection_name = "_" + kbId.replace('-', '_')
        kbName = collection_mapping[collection_name]
    except KeyError:
        return {"code": 200, "message": "知识库不存在"}
    except Exception as e:
        # 处理其他可能的异常
        return {"code": 500, "message": f"发生未知错误: {str(e)}"}

    #首先判断所给kbId对应的知识库是否存在
    if not utility.has_collection(collection_name):
        return {"code": 400, "message": f"知识库不存在"}

    #进行删除知识库操作
    try:
        utility.drop_collection(collection_name)
        #Delete_local_directory
        kbName = collection_mapping[collection_name]
        Knowledge_base_dir = os.path.join(base_dir, kbName)
        shutil.rmtree(Knowledge_base_dir)
        collection_mapping.pop(collection_name)
        save_mapping(collection_mapping_dir,collection_mapping)
        return {"code": 200, "message": "知识库删除成功"}

    except Exception as e:
        return {"code": 500, "message": f"知识库删除失败: {str(e)}"}


@router.post("/delete_knowledge_base_document")
async def delete_files_from_knowledgebase(kbId: str = Body(..., embed=True), files: list[str] = Body(...)):
    """
    根据上传的 fileid 列表删除指定知识库中的对应实体
    :param kbId: 知识库 ID（Milvus collection 名称）
    :param files: 上传的 fileid 列表
    """
    try:
        # 验证输入参数
        if not kbId:
            return {"code": 400, "message": "知识库 ID (kbId) 不能为空"}
        if not files:
            return {"code": 200, "message": "没有需要删除的文件"}

        try:
            # 载入映射表
            collection_mapping = load_mapping(collection_mapping_dir)
            collection_name = "_" + kbId.replace('-', '_')
            kbName = collection_mapping[collection_name]
        except KeyError:
            return {"code": 200, "message": "知识库不存在"}
        except Exception as e:
            # 处理其他可能的异常
            return {"code": 500, "message": f"发生未知错误: {str(e)}"}

        # 首先判断所给 kbId 对应的知识库是否存在
        if not utility.has_collection(collection_name):
            return {"code": 400, "message": f"知识库不存在"}

        collection = Collection(collection_name)

        # 构造 Milvus 查询表达式，根据 fileid 查询对应的 filename
        expr = f"fileid in {files}"
        results = collection.query(
            expr=expr,
            output_fields=["filename"]
        )
        filenames = list(set([result["filename"] for result in results]))

        if not filenames:
            return {"code": 200, "message": "所有 fileid 均未找到对应的文件名"}

        # 构造 Milvus 删除表达式
        delete_expr = f"filename in {filenames}"
        logger.info(f"构造的删除表达式: {delete_expr}")

        # 执行知识库删除操作
        try:
            collection.delete(delete_expr)
            logger.info(f"成功删除知识库 {kbName} 指定文档数据")
        except MilvusException as e:
            logger.error(f"Milvus 删除操作失败: {str(e)}")
            return {"code": 500, "message": f"Milvus 删除操作失败: {str(e)}"}

        # 删除本地文件
        local_deleted = 0
        local_failed = []
        #unique_filenames = set(filenames)  # 去除重复的文件名

        for filename in filenames:
            file_path = os.path.join(base_dir, kbName, filename)
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    local_deleted += 1
                except OSError as e:
                    logger.error(f"删除本地文件失败: {file_path} - {str(e)}")
                    local_failed.append({"filename": filename, "error": str(e)})
            else:
                logger.info(f"文件 {file_path} 不存在，跳过删除操作")

        return {
            "code": 200,
            "message": "文件删除成功",
            "data": {
                "local_deleted": local_deleted,
                "local_failed": local_failed
            }
        }

    except Exception as e:
        logger.error(f"删除过程中发生错误: {str(e)}")
        return {"code": 500, "message": f"服务器内部错误: {str(e)}"}
    
"""
@router.post("/delete_knowledge_base_document")
async def delete_files_from_knowledgebase(kbId: str = Body(..., embed=True), files: List[UploadFile] = File(...)):
    
    #根据上传的文件列表删除指定知识库中的对应实体
    #:param kbId: 知识库ID（Milvus collection名称）
    #:param files: 上传的文件列表（FastAPI UploadFile类型）
    
    try:
        # 验证输入参数
        if not kbId:
            raise ValueError("知识库ID(kbId)不能为空")
        if not files:
            return {"code": 200, "message": "没有需要删除的文件"}

        # 提取所有需要删除的文件名
        filenames = [file.filename for file in files if file.filename]
        if not filenames:
            return {"code": 200, "message": "所有文件均缺少有效文件名"}

        # 构造Milvus查询表达式,expr表达式用列表,用元组会报错无法解析expr表达式
        expr = f"filename in {filenames}"
        logger.info(f"构造的删除表达式: {expr}")

        # 载入映射表
        collection_mapping = load_mapping(collection_mapping_dir)
        collection_name = "_" + kbId.replace('-', '_')
        kbName = collection_mapping[collection_name]

        # 首先判断所给kbId对应的知识库是否存在
        if not utility.has_collection(collection_name):
            return {"code": 400, "message": f"知识库不存在"}

        # 执行知识库删除操作
        try:
            collection = Collection(collection_name)
            collection.delete(expr)
            logger.info(f"成功删除知识库{kbName}指定文档数据")
        except MilvusException as e:
            logger.error(f"Milvus删除操作失败: {str(e)}")
            return {"code": 500, "message": f"Milvus删除操作失败: {str(e)}"}

        # 删除本地文件
        local_deleted = 0
        local_failed = []

        for filename in filenames:
            file_path = os.path.join(base_dir, kbName, filename)
            try:
                os.remove(file_path)
                local_deleted += 1
            except OSError as e:
                logger.error(f"删除本地文件失败: {file_path} - {str(e)}")
                local_failed.append({"filename": filename, "error": str(e)})

        return {
            "code": 200,
            "message": "文件删除成功",
            "data": {
                "local_deleted": local_deleted,
                "local_failed": local_failed
            }
        }

    except Exception as e:
        logger.error(f"删除过程中发生错误: {str(e)}")
        return {"code": 500, "message": f"服务器内部错误: {str(e)}"}
"""
    
"""
# 检索接口
@router.get("/search_knowledge_base")
def search_knowledge_base(query: str, col_name: str, limit: int = 3):
    if not utility.has_collection(col_name):
        return {"message": f"知识库 {col_name} 不存在."}
    try:
        col = Collection(col_name)
        query_embedding = get_text_embeddings([query])[0]
        search_params = {"metric_type": "COSINE", "params": {}}
        res = col.search(
            [query_embedding],
            anns_field="vector",
            limit=limit,
            output_fields=["text", "filename"],
            param=search_params
        )[0]
        results = []
        for hit in res:
            text = hit.get("text")
            filename = hit.get("filename")
            results.append({"text": text, "filename": filename})
        return {"results": results}
    except Exception as e:
        return {"message": f"知识库 {col_name} 检索失败. 报错如下: {str(e)}"}
"""


# 允许的文件类型
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}

def allowed_file(filename: str) -> bool:
    """检查文件扩展名是否合法"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

"""
#pdfplumber解析
async def parse_file(file_path: str, extension: str) -> list[str]:
    #解析不同格式文件为文本列表
    texts = []
    try:
        if extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                texts = [f.read()]
        
        elif extension == 'pdf':
            # 使用pdfplumber解析（更精准保留文本结构）
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    texts.append(page.extract_text())
        
        elif extension == 'docx':
            doc = Document(file_path)
            texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        elif extension == 'doc':
            # 注意：python-docx不支持.doc格式，需要调用其他库或转换
            raise NotImplementedError("DOC格式需要先转换为DOCX")

    except Exception as e:
        logger.error(f"文件解析失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    return [text for text in texts if text.strip()]
"""

"""
import fitz  # PyMuPDF解析
async def parse_file(file_path: str, extension: str) -> list[str]:
    #解析不同格式文件为文本列表
    texts = []
    try:
        if extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                texts = [f.read()]

        elif extension == 'pdf':
            # 使用PyMuPDF解析
            doc = fitz.open(file_path)
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text()
                texts.append(text)
            doc.close()

        elif extension == 'docx':
            doc = Document(file_path)
            texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

        elif extension == 'doc':
            # 注意：python-docx不支持.doc格式，需要调用其他库或转换
            raise NotImplementedError("DOC格式需要先转换为DOCX")

    except Exception as e:
        logger.error(f"文件解析失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    return [text for text in texts if text.strip()]
"""

import io
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage

async def parse_file(file_path: str, extension: str):
    """解析不同格式文件为文本列表"""
    parse_contents = []
    texts = []
    try:
        if extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                parse_contents = [f.read()]
            texts = parse_contents            
        
        elif extension == 'pdf' or extension == 'docx':
            
            #使用file_mineru函数进行解析
            #md_list = process_pdf_file(file_path)
            
            data = {
                "file_paths": [file_path],
                "output_dir": output_dir
            }
            
            try:
                response = requests.post(url=url,json=data)
                md_list = response.json()[0]
                md_file_path, md_content = md_list['output_file'], md_list['md_content']
                parse_contents = [md_content]
                logger.info("md文件解析成功")
                #再对md文件进行分块处理
                #file_name, merged_chunks = process_markdown_file(md_file_path)
                #parser = EnhancedMarkdownParser()
                #merged_chunks = parser.parse(md_file_path) 
                merged_chunks = process_markdown_file(md_file_path)
                logger.info("md文件完成分块处理")
                texts = merged_chunks
            
            except Exception as e:
                logger.error(f"文件解析失败: {str(e)}")
                raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")
        
        elif extension == 'doc':
            # 注意：python-docx 不支持.doc 格式，需要调用其他库或转换
            raise NotImplementedError("DOC 格式需要先转换为 DOCX")
            
        """
        elif extension == 'pdf':
            # 使用 pdfminer.six 解析
            rsrcmgr = PDFResourceManager()
            retstr = io.StringIO()
            laparams = LAParams()
            device = TextConverter(rsrcmgr, retstr, laparams=laparams)
            with open(file_path, 'rb') as fp:
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                for page in PDFPage.get_pages(fp):
                    interpreter.process_page(page)
            text = retstr.getvalue()
            device.close()
            retstr.close()
            texts = [text]

        elif extension == 'docx':
            doc = Document(file_path)
            texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

        """
        

    except Exception as e:
        logger.error(f"文件解析失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    return [parse_content for parse_content in parse_contents if parse_content.strip()],[text for text in texts if text.strip()]

#密集检索
def dense_search(col, query_dense_embedding, limit):
    search_params = {"metric_type": "COSINE", "params": {}}
    res = col.search(
        [query_dense_embedding],
        anns_field="dense_vector",
        limit=limit,
        output_fields=["text"],
        param=search_params
    )[0]
    return [hit.get("text") for hit in res]

#稀疏检索
def sparse_search(col, query_sparse_embedding, limit):
    search_params = {
        "metric_type": "IP",
        "params": {},
    }
    res = col.search(
        [query_sparse_embedding],
        anns_field="sparse_vector",
        limit=limit,
        output_fields=["text"],
        param=search_params
    )[0]
    return [hit.get("text") for hit in res]

#混合检索
def hybrid_search(
        col,
        query_dense_embedding,
        query_sparse_embedding,
        sparse_weight,
        dense_weight,
        limit
):
    dense_search_params = {"metric_type": "COSINE", "params": {}}
    dense_req = AnnSearchRequest(
        [query_dense_embedding], "dense_vector", dense_search_params, limit=limit
    )
    sparse_search_params = {"metric_type": "IP", "params": {"drop_ratio_build": 0.2,"drop_ratio_search": 0.2}}
    sparse_req = AnnSearchRequest(
        [query_sparse_embedding], "sparse_vector", sparse_search_params, limit=limit
    )
    rerank = WeightedRanker(sparse_weight, dense_weight)
    #rerank = RRFRanker()
    res = col.hybrid_search(
        [sparse_req, dense_req], rerank=rerank, limit=limit, output_fields=["text"]
    )[0]
    return [hit.get("text") for hit in res]


"""""查找与当前上传文件相似度最高的文件接口"""
@router.post("/search_by_file")
async def search_by_file(kbId: str = Body(...),file: UploadFile = File(...)):

    # 1. 验证文件类型
    if not allowed_file(file.filename):
        return {"code": 400, "message": "仅支持PDF/DOC/DOCX/TXT文件"}

    try:
        # 载入映射表
        collection_mapping = load_mapping(collection_mapping_dir)
        collection_name = "_" + kbId.replace('-', '_')
        kbName = collection_mapping[collection_name]
    except KeyError:
        return {"code": 200, "message": "知识库不存在"}
    except Exception as e:
        # 处理其他可能的异常
        return {"code": 500, "message": f"发生未知错误: {str(e)}"}

    # 2. 保存本地文件
    file_ext = file.filename.rsplit('.', 1)[1].lower()
    temp_path = os.path.join(os.path.join(base_dir, kbName),file.filename)
    
    #if os.path.exists(temp_path):
        #return{"code":500, "message":"系统检测到相同文件名，请核实上传文件"}
    
    try:
        contents = await file.read()
        with open(temp_path, 'wb') as f:
            f.write(contents)
    
    except Exception as e:
        logger.error(f"文件保存本地失败: {str(e)}")
        return {"code": 500, "message": "文件处理失败"}
    
    # 3. 解析文件内容
    try:
        parse_contents,query_texts = await parse_file(temp_path, file_ext)
        logger.info(f"解析内容如下: {query_texts}")
        logger.info(f"解析出{len(query_texts)}段文本")
        if not query_texts:
            os.remove(temp_path)
            return {"code": 400, "message": "文件内容为空或无法解析"}
    except Exception as e:
        os.remove(temp_path)
        return {"code": 500, "message": f"内容解析失败: {str(e)}"}
    
    # 4. 待查询文本分块
    #chunked_texts = split_text_into_chunks(query_texts, 512)
    #logger.info(f"分块后的待查询文本列表长度为:{len(chunked_texts)}")
    #logger.info(f"分块后的待查询文本列表:{chunked_texts}")

    # 5. 执行相似性搜索
    try:
        if not utility.has_collection(collection_name):
            os.remove(temp_path)
            return {"code": 400, "message": "知识库不存在"}
        col = Collection(collection_name)
        #query_embedding = emb_text(query_texts,is_query=False)
        
        #混合检索
        query_embedding = generate_embeddings(query_texts)
        
        #dense_search_params = {"metric_type": "COSINE", "params": {"nprobe": 64}}
        dense_search_params = {"metric_type": "IP", "params": {"nprobe": 64}}
        dense_req = AnnSearchRequest(
            [query_embedding["dense"][0]], "dense_vector", dense_search_params, limit=50
        )
        
        #sparse_search_params = {"metric_type": "IP", "params": {"nprobe": 16}}
        sparse_search_params = {"metric_type": "IP", "params": {"top_k": 100}}
        
        sparse_req = AnnSearchRequest(
            [query_embedding["sparse"]._getrow(0)], "sparse_vector", sparse_search_params, limit=50
        )
        
        """
        sparse_search_params = {"metric_type": "BM25"}
        sparse_req = AnnSearchRequest(
            [query_texts], "sparse_bm25", sparse_search_params, limit=10
        )
        """
        #sparse_weight=0.6
        #dense_weight=0.4
        #rerank = WeightedRanker(sparse_weight, dense_weight)
        rerank = RRFRanker(100)
        results = col.hybrid_search(
            [sparse_req, dense_req], 
            rerank=rerank, 
            limit=10, 
            output_fields=["text", "filename","fileid"]
        )[0]
        
        """
        # 搜索参数（可根据需求调整）
        search_params = {
            "metric_type": "COSINE",
            "params": {"nprobe": 32}  # 提高搜索精度
        }
        
        results = col.search(
            data=query_embedding,
            anns_field="vector",
            param=search_params,
            limit=10,
            output_fields=["text", "filename","fileid"]
        )[0]
        
        """
        # 格式化结果
        formatted_results = []
        for hit in results:
            formatted_results.append({
                "text": hit.entity.get("text"),
                "filename": hit.entity.get("filename"),
                "fileid": hit.entity.get("fileid"),
                "similarity": round(hit.score, 4)  # 保留4位小数
            })
        logger.info(f"搜索成功相似文件结果如下:{formatted_results}")
        
        """
        documents = []
        
        for result in formatted_results:
            documents.append(result["text"])
        
        bge_rf = BGERerankFunction(model_name='bge-reranker-v2-m3',device='cuda:0')
        rerank_results = bge_rf(
            query=query_texts,
            documents=documents,
            top_k=3
        )
        
        #return {"query_text": combined_text[:200]+"...", "results": formatted_results}
        #return {"docs": formatted_results[0]["text"],"relevantScore": formatted_results[0]["similarity"],"relFile": formatted_results[0]["filename"]}

        final_result = []
        for i, result in enumerate(rerank_results):
            original_result = next((res for res in formatted_results if res["text"] == result.text), None)
            if original_result:
                final_result.append({
                    "docs": original_result["text"],
                    "relevantScore": result.score,
                    "relFile": original_result["filename"],
                    "relFileId": original_result["fileid"]
                })
        
        """
        
        final_result = []
        for result in formatted_results:
            final_result.append({"docs": result["text"],"relevantScore": result["similarity"],"relFile": result["filename"],"relFileId": result["fileid"]})
          
        
        return {"code": 200, "files":final_result,"parse_content":parse_contents}

    except Exception as e:
        os.remove(temp_path)
        logger.error(f"搜索失败: {str(e)}")
        return {"code": 500, "message": f"搜索失败: {str(e)}"}



"""确认文件入库接口(将文件存入知识库)"""
# 入库接口
@router.post("/upload_file_confirm")
async def upload_file_confirm(kbId: str = Body(...),fileId: str = Body(...),fileName: str = Body(...),file: UploadFile = File(...),updateContent: Optional[str] = Body(None)):

     # 1. 检查文件类型
    if not allowed_file(fileName):
        return {"code": 400, "message": "仅支持PDF/DOC/DOCX/TXT文件"}
    
    try:
        # 载入映射表
        collection_mapping = load_mapping(collection_mapping_dir)
        collection_name = "_" + kbId.replace('-', '_')
        kbName = collection_mapping[collection_name]
    except KeyError:
        return {"code": 200, "message": "知识库不存在"}
    except Exception as e:
        # 处理其他可能的异常
        return {"code": 500, "message": f"发生未知错误: {str(e)}"}

    # 2. 获取上传文件路径以及
    file_ext = fileName.rsplit('.', 1)[1].lower()
    file_path = os.path.join(os.path.join(base_dir, kbName),fileName)
    
    """
     # 3. 解析文件内容
    try:
        texts = await parse_file(file_path, file_ext)
        #logger.info(f"解析内容如下: {texts}")
        logger.info(f"解析出{len(texts)}段文本")
    
    except NotImplementedError as e:
        os.remove(file_path)  # 删除无效文件
        return {"code": 400, "message": f"报错信息: {str(e)}"}
    """

    # 3. 读取已经解析好的对应文本内容
    if file_ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            parse_contents = [f.read()]
        texts = parse_contents
        
    else:
        #读取当前文件解析后的 md文件
        filename = fileName.rsplit('.', 1)[0]
        md_file_path = os.path.join(output_dir, filename, 'txt', filename + '.md')

        #再对其md文件进行分块处理
        #_ , merged_chunks = process_markdown_file(md_file_path)
        #parser = EnhancedMarkdownParser()
        #merged_chunks = parser.parse(md_file_path) 
        merged_chunks = process_markdown_file(md_file_path)
        texts = merged_chunks

    
    # 4. 检查集合是否存在
    if not utility.has_collection(collection_name):
        os.remove(file_path)
        return {"code": 400, "message": "知识库不存在"}


    # 5. 生成嵌入并插入Milvus    
    try:
        col = Collection(collection_name)
        logger.info("-----开始文本向量化-----")
        #先分块
        #chunked_texts = split_text_into_chunks(texts, 512)
        #logger.info(f"分块后的文本列表:{chunked_texts}")
        #embeddings = emb_text(texts,is_query=False)
        embeddings = generate_embeddings(texts)
        #logger.info(f"向量化后的结果:{embeddings}")
        logger.info("-----完成文本向量化-----")
        logger.info(f"分块处理:{texts}")
        # 将每个元素转换为列表的列表格式
        entities = [
            texts,  #  texts 为字符串列表
            embeddings['sparse'],
            embeddings['dense'],
            [fileName] * len(texts),  # 重复文件名以匹配文本数量
            [fileId] * len(texts)  # 重复文件ID以匹配文本数量
        ]

        """
        entities = [
            texts,  #  texts 为字符串列表
            embeddings,
            [fileName] * len(texts),  # 重复文件名以匹配文本数量
            [fileId] * len(texts)  # 重复文件ID以匹配文本数量
        ]
        """
        col.insert(entities)
        col.flush()  # 确保数据持久化
       
        logger.info(f"文档已成功上传至知识库: --{kbName}-- ")
        return { "code": 200,"message": "文件向量化并成功入库"}

    
    except Exception as e:
        os.remove(file_path)
        logger.error(f"Milvus操作失败: {str(e)}")
        return {"code": 500, "message": f"上传失败: {str(e)}"}

"""知识检索接口（单纯检索知识库，无模型问答）"""
@router.post("/retrieve_knowledge")
def retrieval(kbId: Optional[str] = Body(None),query: str = Body(...),limit: int = Body(default=20)):
    """
    # 查询重构
    messages = [
        {
            'role': 'system',
            'content': '你现在的身份是医疗助手'
        },
        {
            'role': 'user',
            'content': f'请根据以下文本生成一段背景知识，文本为：{query}'
        },
    ]
    
    response = client.chat.completions.create(
        model='qwen-max',  # 此处以qwen-max为例，可按需更换模型名称。模型列表：https://help.aliyun.com/zh/model-studio/getting-started/models
        messages=messages,
        stream=False,
        temperature=0
    )
    
    query = response.choices[0].message.content
    """
    #查询query文本向量化
    #query_embedding = emb_text(query,is_query=True)
    query_embedding = generate_embeddings(query)
    
    #首先判断kbId字段是否传入，若传入则进行对应知识库检索，否则进行所有知识库检索
    #所有知识库检索
    if kbId is None or kbId == '' :
        try:
            # 获取所有 Collection 名称
            collection_names = utility.list_collections()
            logger.info(f"所有 Collection 名称: {collection_names}")
            # 初始化结果列表
            all_results = []
            # 遍历每个 Collection
            for collection_name in collection_names:

                #载入映射表
                collection_mapping = load_mapping(collection_mapping_dir)
                kbName = collection_mapping[collection_name]

                # 检索相关文件
                col = Collection(collection_name)

                
                #混合检索
                #dense_search_params = {"metric_type": "COSINE", "params": {"nprobe": 64}}
                dense_search_params = {"metric_type": "IP", "params": {"nprobe": 64}}
                dense_req = AnnSearchRequest(
                    [query_embedding["dense"][0]], "dense_vector", dense_search_params, limit=(limit+50)
                )
                #sparse_search_params = {"metric_type": "IP", "params": {"nprobe": 16}}
                sparse_search_params = {"metric_type": "IP", "params": {"top_k": 100}}
                """
                sparse_search_params = {"metric_type": "BM25"}
                sparse_req = AnnSearchRequest(
                    [query], "sparse_bm25", sparse_search_params, limit=limit
                )
                """
                sparse_req = AnnSearchRequest(
                    [query_embedding["sparse"]._getrow(0)], "sparse_vector", sparse_search_params, limit=(limit+50)
                )
                
                #sparse_weight=0.6
                #dense_weight=0.4
                #rerank = WeightedRanker(sparse_weight, dense_weight)
                rerank = RRFRanker(100)
                results = col.hybrid_search(
                        [sparse_req, dense_req], 
                        rerank=rerank, 
                        limit=limit, 
                        output_fields=["text", "filename","fileid"]
                )[0]
                
                """
                # 搜索参数（可根据需求调整）
                search_params = {
                    "metric_type": "COSINE",
                    "params": {"nprobe": 32}  # 提高搜索精度
                }


                results = col.search(
                    data=query_embedding,
                    anns_field="vector",
                    param=search_params,
                    limit=limit,
                    output_fields=["text", "filename","fileid"]
                )[0]
                """

                # 格式化结果
                formatted_results = []
                for hit in results:
                    formatted_results.append({
                        "text": hit.entity.get("text"),
                        "filename": hit.entity.get("filename"),
                        "fileid": hit.entity.get("fileid"),
                        "filepath": os.path.join(base_dir, kbName, hit.entity.get("filename")),
                        "similarity": round(hit.score, 4)  # 保留4位小数
                    })
                all_results.append(formatted_results)

            # 合并所有结果
            merged_results = []
            for formatted_results in all_results:
                merged_results.extend(formatted_results)
            
            # 根据 similarity 进行降序排序
            sorted_results = sorted(merged_results, key=lambda x: x["similarity"], reverse=True)

            # 取前limit的相关文档
            top_k_results = sorted_results[:limit]
            #logger.info(f"检索成功，相关文件结果如下:{top_k_results}")
            
            """
            final_result = []
            i = 1
            for result in top_k_results:
                final_result.append({"num": i,"content": result["text"],"fileName": result["filename"],"fileId": result["fileid"] ,"score": result["similarity"],"filePath": result["filepath"]})
                i+=1
            #return {"content": top_three_results[0]["text"],"relevantScore": top_three_results[0]["similarity"],"relFile": top_three_results[0]["filename"]}
            
            return {"code": 200,"files":final_result}

        except Exception as e:
            logger.error(f"搜索失败: {str(e)}")
            return {"code": 500, "message": f"检索失败: {str(e)}"}
        """
        
            # 使用 BGERerankFunction 进行重排
            documents = [res["text"] for res in top_k_results]
            bge_rf = BGERerankFunction(model_name='bge-reranker-large', device='cuda:0')
            rerank_results = bge_rf(
                query=query,
                documents=documents,
                top_k=3
            )

            # 根据重排结果重新构建最终结果
            final_result = []
            for i, rerank_result in enumerate(rerank_results):
                for res in top_k_results:
                    if res["text"] == rerank_result.text:
                        final_result.append({
                            "num": i + 1,
                            "content": res["text"],
                            "fileName": res["filename"],
                            "fileId": res["fileid"],
                            "score": rerank_result.score,
                            "filePath": res["filepath"]
                        })
            
            logger.info(f"检索成功，相关文件结果如下:{final_result}")
            
            return {"code": 200, "files": final_result}

        except Exception as e:
            logger.error(f"搜索失败: {str(e)}")
            return {"code": 500, "message": f"检索失败: {str(e)}"}
        
    else :

        #指定知识库检索
        
        try:
        # 载入映射表
            collection_mapping = load_mapping(collection_mapping_dir)
            collection_name = "_" + kbId.replace('-', '_')
            kbName = collection_mapping[collection_name]
        except KeyError:
            return {"code": 200, "message": "知识库不存在"}
        except Exception as e:
            # 处理其他可能的异常
            return {"code": 500, "message": f"发生未知错误: {str(e)}"}

        # 检索相关文件
        try:
            if not utility.has_collection(collection_name):
                return {"code": 400, "message": "知识库不存在"}
            col = Collection(collection_name)
            
            #spare_result = sparse_search(col,query_embedding["sparse"]._getrow(0),limit=limit)
            #dense_result = dense_search(col,query_embedding["dense"][0],limit=limit)
            #logger.info(f'------sparse_result :{spare_result}')
            #logger.info(f'------dense_result :{dense_result}')
            #混合检索
            #dense_search_params = {"metric_type": "COSINE", "params": {"nprobe": 64}}
            
            dense_search_params = {"metric_type": "IP", "params": {"nprobe": 64}}
            dense_req = AnnSearchRequest(
                [query_embedding["dense"][0]], "dense_vector", dense_search_params, limit=(limit+50)
            )
            #sparse_search_params = {"metric_type": "IP", "params": {"nlist": 128}}
            sparse_search_params = {"metric_type": "IP", "params": {"top_k": 100}}
            
            sparse_req = AnnSearchRequest(
                [query_embedding["sparse"]._getrow(0)], "sparse_vector", sparse_search_params, limit=(limit+50)
            )
            """
            sparse_search_params = {"metric_type": "BM25"}
            sparse_req = AnnSearchRequest(
                [query], "sparse_bm25", sparse_search_params, limit=limit
            )
            """
            #sparse_weight=0.6
            #dense_weight=0.4
            #rerank = WeightedRanker(sparse_weight,dense_weight)
            rerank = RRFRanker(100)
            results = col.hybrid_search(
                    [sparse_req, dense_req], 
                    rerank=rerank, 
                    limit=limit, 
                    output_fields=["text", "filename","fileid"]
            )[0]
            #query_embedding = emb_text(query,is_query=True)
            """
            # 搜索参数（可根据需求调整）
            search_params = {
                "metric_type": "COSINE",
                "params": {"nprobe": 32}  # 提高搜索精度
            }
        
            results = col.search(
                data=query_embedding,
                anns_field="vector",
                param=search_params,
                limit=limit,
                output_fields=["text", "filename","fileid"]
            )[0]
            """

            
            formatted_results = []

            for hit in results:
                formatted_results.append({
                    "text": hit.entity.get("text"),
                    "filename": hit.entity.get("filename"),
                    "fileid": hit.entity.get("fileid"),
                    "filepath": os.path.join(base_dir, kbName, hit.entity.get("filename")),
                    "similarity": round(hit.score, 4)  # 保留4位小数
                })
            
            #logger.info(f"检索成功，相关文件结果如下:{formatted_results}")
                
            # 使用 BGERerankFunction 进行重排
            documents = [res["text"] for res in formatted_results]
            bge_rf = BGERerankFunction(model_name='bge-reranker-large', device='cuda:0')
            rerank_results = bge_rf(
                query=query,
                documents=documents,
                top_k=3
            )

            # 根据重排结果重新构建最终结果
            final_result = []
            for i, rerank_result in enumerate(rerank_results):
                for res in formatted_results:
                    if res["text"] == rerank_result.text:
                        final_result.append({
                            "num": i + 1,
                            "content": res["text"],
                            "fileName": res["filename"],
                            "fileId": res["fileid"],
                            "score": rerank_result.score,
                            "filePath": res["filepath"]
                        })

            logger.info(f"检索成功，相关文件结果如下:{final_result}")

            return {"code": 200, "files": final_result}

        except Exception as e:
            logger.error(f"搜索失败: {str(e)}")
            return {"code": 500, "message": f"检索失败: {str(e)}"}
            
            """
            #logger.info(f'结果:{results}')
            # 格式化结果
            formatted_results = []

            for hit in results:
                formatted_results.append({
                    "text": hit.entity.get("text"),
                    "filename": hit.entity.get("filename"),
                    "fileid": hit.entity.get("fileid"),
                    "filepath": os.path.join(base_dir, kbName, hit.entity.get("filename")),
                    "similarity": round(hit.score, 4)  # 保留4位小数
                })
            
            logger.info(f"检索成功，相关文件结果如下:{formatted_results}")

            final_result = []
            i = 1
            for result in formatted_results:
                final_result.append({"num": i,"content": result["text"],"fileName": result["filename"],"fileId": result["fileid"] ,"score": result["similarity"],"filePath": result["filepath"]})
                i+=1
        
            return {"code": 200,"files":final_result}
        
        except Exception as e:
            logger.error(f"搜索失败: {str(e)}")
            return {"code": 500, "message": f"检索失败: {str(e)}"}
        """

# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="localhost", port=7070)